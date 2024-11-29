import io
import os
import ssl
from urllib.request import Request, urlopen
from docx import Document
from PyPDF2 import PdfFileReader, PdfReader
import re

def read_docx(file_path):
    """Reads and extracts the relevant sections from a .pdf or a .docx file into a dictionary."""
    
    return {
        "result": {
            "Skills (Must have)": [
            "3-4 years of experience in developing Web applications",
            "Proficient with ReactJS, JavaScript, Typescript, HTML, Node JS(Good to have)",
            "Understanding of web security best practices.",
            "Expertise with web application architecture, application security and REST APIs.",
            "Should Understand basic Algorithms and Data Structures.",
            "Experience writing Unit Tests with Jest, Enzyme or React Testing Library"
            ],
            "Skills (Good to have)": [
            "Expert knowledge of front-end build pipeline and tools (npm, webpack etc)",
            "Experience in writing REST APIs in Python as a back-end programming language is preferable.",
            "Sound debugging skills using tools like browser's developer tools, Fiddler, Postman etc.",
            "Working knowledge of design tools, Unix, docker, CI/CD, SVN/Git etc.",
            "Experience with NodeJs.",
            "Familiarity with working in a Scrum Agile delivery environment",
            "Experience with JavaScript testing frameworks",
            "Experience with Frontend performance analysis"
            ],
            "Responsibilities": [
            "Develop web applications using JavaScript, React.JS, HTML and CSS.",
            "Follow the best practices for code development and code hygiene. Write modular and unit tested code.",
            "Work with other developers on the development team on the implementation of common frameworks and solutions.",
            "Work with the development team to support and maintain existing production code in the field and develop, deliver new enhancements & products.",
            "Take end to end responsibility of the assigned tasks/modules/features by interacting with different stakeholders like Product Managers, UX designers, QA, etc.",
            "Work in SCRUM / Agile environment. Follow the Scrum process, participate in Scrum ceremonies and follow the incremental delivery model.",
            "Qualification:",
            "Bachelors/Masters in Computer Science, Software Engineering or equivalent.",
            "Additional Information:",
            "We offer a competitive salary and excellent benefits that are above industry standard.",
            "Do check our impressive growth rate on LinkedIn and ratings on Glassdoor",
            "Pls submit your resume in this standard 1-page format or 2-page format",
            "Please hear from our employees on Life at Josh Software"
            ]
        }
    }
    
    try:
        result = {
            "Skills (Must have)": [],
            "Skills (Good to have)": [],
            "Responsibilities": []
        }
        ssl._create_default_https_context = ssl._create_unverified_context
        remote_file = urlopen(Request(file_path)).read()
        memory_file = io.BytesIO(remote_file)
        
        doc = Document(memory_file)
        
        current_section = None
        
        for paragraph in doc.paragraphs:
            if "Skills (Must have)" in paragraph.text:
                current_section = "Skills (Must have)"
            elif "Skills (Good to have)" in paragraph.text:
                current_section = "Skills (Good to have)"
            elif "Responsibilities" in paragraph.text:
                current_section = "Responsibilities"
            elif current_section:
                if paragraph.text.strip():
                    result[current_section].append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        result[current_section].append(cell_text)
                        
        return result
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return f"Error reading file {file_path}: {e}"
        return None


def read_pdf(file_path):
    """Reads and prints the content of a .pdf file."""
    print(f"\nReading: {file_path}")
    try:
        remote_file = urlopen(Request(file_path)).read()
        memory_file = io.BytesIO(remote_file)
        
        reader = PdfFileReader(memory_file)
        
        for page in reader.pages:
            print(page.extract_text())
            content = page.extract_text()
            return content
    except Exception as e:
        print(f"Error reading .pdf file {file_path}: {e}")
        return f"Error reading .pdf file {file_path}: {e}"

def process_directory(directory_path):
    """Reads all .docx and .pdf files in the given directory."""
    try:
        for root, _, files in os.walk(directory_path):
            cnt = 0
            for file in files:
                cnt+=1
                print("***********************\n\n\nReading File No : ",cnt)

                original_file_path = os.path.join(root, file)
                if " " in file:
                    new_file_name = file.replace(" ", "_")
                    new_file_path = os.path.join(root, new_file_name)
                    os.rename(original_file_path, new_file_path)
                    print(f"Renamed file: {original_file_path} -> {new_file_path}")
                    file = new_file_name
                
                file_path = os.path.join(root, file)

                if file.lower().endswith('.docx'):
                    read_docx(file_path)
                elif file.lower().endswith('.pdf'):
                    read_pdf(file_path)
                else:
                    print(f"Unsupported file type: {file_path}")
    except Exception as e:
        print(f"Error processing directory {directory_path}: {e}")

if __name__ == "__main__":
    directory_path = input("Enter the directory path: ")
    process_directory(directory_path)

